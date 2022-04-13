from flask import Flask, send_from_directory, send_file, safe_join, abort
from flask import request, jsonify
import json
from docxtpl import DocxTemplate
from docx import Document
from flask_cors import CORS
import inspect
import os
import re
from werkzeug.utils import secure_filename
from flask_swagger_ui import get_swaggerui_blueprint
# from routes import request_api
# import win32com.client
# import pythoncom

app = Flask(__name__)

### swagger specific ###
SWAGGER_URL = '/docs'
API_URL = '/static/swagger.json'
SWAGGERUI_BLUEPRINT = get_swaggerui_blueprint(
    SWAGGER_URL,
    API_URL,
    config={
        'app_name': "Seans-Python-Flask-REST-Boilerplate"
    }
)
app.register_blueprint(SWAGGERUI_BLUEPRINT, url_prefix=SWAGGER_URL)
### end swagger specific ###
# app.register_blueprint(request_api.get_blueprint())


# app.config["DEBUG"] = True
UPLOAD_FOLDER = '/var/www/docxapi/uploads'
# UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
CORS(app)

# @app.route('/<name>')
# def index(name):
#     return '<h1>Hello </h1>'.format(name)


@app.route('/', methods=['GET'])
def home():
    return "<h1>Distant Reading Archive</h1><p>This site is a prototype API for distant reading of science fiction novels.</p>"


@app.route('/static/<path:path>')
def send_static(path):
    return send_from_directory('static', path)


def update_toc(docx_file):
    word = win32com.client.DispatchEx(
        "Word.Application", pythoncom.CoInitialize())
    doc = word.Documents.Open(docx_file)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()

# A route to return all of the available entries in our catalog.


@app.route('/api/v1/resources/firstapi', methods=['GET', 'POST'])
def first_api():
    if request.method == 'POST':
        data = []
        output_file_path = 'result.docx'
        # template_file_path = 'Retelit_Products.docx'
        if "document" in request.files:
            document = request.files["document"]
            template_file_path = document.filename
            document.save(os.path.join(
                app.config['UPLOAD_FOLDER'], document.filename))
            # document.save(os.path.join(document.filename))
            # print("path: "+app.config['UPLOAD_FOLDER']+'/'+template_file_path)
            template_document = Document(
                app.config['UPLOAD_FOLDER']+'/'+template_file_path)

            replacesImageName = []
            files = request.files.getlist("replace_images[]")
            for file in files:
                print(file.filename)
                replacesImageName.append(
                    app.config['UPLOAD_FOLDER']+'/'+file.filename)
                file.save(os.path.join(
                    app.config['UPLOAD_FOLDER'], file.filename))

            findsImageName = []
            files = request.files.getlist("find_images[]")
            for file in files:
                print(file.filename)
                findsImageName.append(
                    app.config['UPLOAD_FOLDER']+'/'+file.filename)
                file.save(os.path.join(
                    app.config['UPLOAD_FOLDER'], file.filename))

            find_text = request.form.getlist('find_text[]')
            replace_text = request.form.getlist('replace_text[]')
            words = []
            if len(find_text) == len(replace_text):
                for i in range(len(replace_text)):
                    prepareWord = {}
                    prepareWord['find'] = find_text[i]
                    prepareWord['replace'] = replace_text[i]
                    words.append(prepareWord)

            words_not_replace = []
            if words:
                for word in words:
                    for paragraph in template_document.paragraphs:
                        if word['find'] in paragraph.text:
                            inline = paragraph.runs
                            for i in range(len(inline)):
                                if word['find'] in inline[i].text:
                                    text = inline[i].text.replace(
                                        word['find'], word['replace'])
                                    inline[i].text = text
                                else:
                                    prepareWord = {}
                                    prepareWord['find'] = word['find']
                                    prepareWord['replace'] = word['replace']
                                    status = True
                                    for rrword in words_not_replace:
                                        if rrword['replace'] == word['replace']:
                                            status = False

                                    if status:
                                        words_not_replace.append(prepareWord)
                                    # paragraph.text = paragraph.text.replace(word['find'], word['replace'])
                            # paragraph.text = paragraph.text.replace(word['find'], word['replace'])
                            # replace_text_in_paragraph(paragraph, word['find'], word['replace'])

                    # for paragraph in template_document.paragraphs:
                    #     print("1 : "+ paragraph.text)
                    #     if word['find'] in paragraph.text:
                    #         paragraph.text = paragraph.text.replace(word['find'], word['replace'])

                    for table in template_document.tables:
                        for col in table.rows:
                            for cell in col.cells:
                                for paragraph in cell.paragraphs:
                                    if word['find'] in paragraph.text:
                                        inline = paragraph.runs
                                        for i in range(len(inline)):
                                            if word['find'] in inline[i].text:
                                                print("0 : " + inline[i].text)
                                                text = inline[i].text.replace(
                                                    word['find'], word['replace'])
                                                inline[i].text = text
                                            else:
                                                prepareWord = {}
                                                prepareWord['find'] = word['find']
                                                prepareWord['replace'] = word['replace']
                                                words_not_replace.append(
                                                    prepareWord)
                                                # paragraph.text = paragraph.text.replace(word['find'], word['replace'])
                                        # replace_text_in_paragraph(paragraph, word['find'], word['replace'])

                    # for table in template_document.tables:
                    #     for row in table.rows:
                    #         for cell in row.cells:
                    #             for paragraph in cell.paragraphs:
                    #                 paragraph.text = paragraph.text.replace(word['find'], word['replace'])
                                    # paragraph.text.bold = False
                                    # replace_text_in_paragraph(paragraph, word['find'], word['replace'])
                print(words_not_replace)
                for rrword in words_not_replace:
                    for paragraph in template_document.paragraphs:
                        if rrword['find'] in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                rrword['find'], rrword['replace'])
                            # print(" etext : "+paragraph.text)

                for rword in words_not_replace:
                    for table in template_document.tables:
                        for col in table.rows:
                            for cell in col.cells:
                                for paragraph in cell.paragraphs:
                                    # print("text : "+paragraph.text)
                                    if rword['find'] in paragraph.text:
                                        paragraph.text = paragraph.text.replace(
                                            rword['find'], rword['replace'])

                for word in words:
                    for section in template_document.sections:
                        for dh in section.first_page_header.paragraphs:
                            if(word['find'] in dh.text):
                                if(dh.text.bold):
                                    word['replace'].bold = True
                                    dh.text = dh.text.replace(
                                        word['find'], word['replace'])

                    for section in template_document.sections:
                        for dfh in section.header.paragraphs:
                            if word['find'] in dfh.text:
                                if(dfh.text.bold):
                                    word['replace'].bold = True
                                    dfh.text = dfh.text.replace(
                                        word['find'], word['replace'])
                    #
                    for section in template_document.sections:
                        for df in section.first_page_footer.paragraphs:
                            if word['find'] in df.text:
                                df.text = df.text.replace(
                                    word['find'], word['replace'])

                    for section in template_document.sections:
                        for dff in section.footer.paragraphs:
                            if word['find'] in dff.text:
                                dff.text = dff.text.replace(
                                    word['find'], word['replace'])

            # template_document.save(output_file_path)
            file_path = os.path.join(
                app.config['UPLOAD_FOLDER'], output_file_path)
            template_document.save(file_path)

            doc = DocxTemplate(
                app.config['UPLOAD_FOLDER']+'/'+output_file_path)
            doc.reset_replacements()
            if len(findsImageName) == len(replacesImageName):
                for i in range(len(findsImageName)):
                    doc.replace_media(findsImageName[i], replacesImageName[i])
            # doc.replace_media('header.jpg','header-replace.jpg')
            # doc.replace_media('main.png','main-replace.png')
            # doc.replace_media('map.png','map-replace.png')
            # return_file = doc
            doc.save(file_path)
            file_path_toc = file_path
            script_dir = os.path.dirname(os.path.abspath(
                inspect.getfile(inspect.currentframe())))
            # print("script_dir : "+file_path_toc)
            # print("script_dir : "+os.path.join(script_dir,file_path_toc))
            # update_toc(os.path.join(script_dir,file_path_toc))

            # words = request.form.getlist('words')
            # return send_from_directory(directory='/', filename=return_file, as_attachment=True)
            # filename = os.path.join(app.root_path, '/', output_file_path)
            # return_data = [os.path.dirname(app.instance_path), output_file_path]
            download_file_url = request.url_root+"download/result.docx"

            # return jsonify(isError= False,
            #                 message= "Success",
            #                 statusCode= 200,
            #                 data= download_file_url), 200

            return jsonify(isError=False,
                           message="Success",
                           statusCode=200,
                           data='result.docx'), 200
            # return send_from_directory(app.config["UPLOAD_FOLDER"], 'result.docx', as_attachment=True)
        else:
            return jsonify(isError=True,
                           message="Document Not Found",
                           statusCode=200,
                           data=[])
    else:
        return jsonify(isError=True,
                       message="Request method not allowed",
                       statusCode=200,
                       data=[])


@app.route('/download/<filename>')
def downloadFile(filename):
    return send_file('uploads/'+filename, as_attachment=True)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


# app.run(debug=True ,port=8080,use_reloader=False)
# loop
if __name__ == "__main__":
    app.run()
#     from waitress import serve
#     serve(debug=True ,port=8080,use_reloader=False)
